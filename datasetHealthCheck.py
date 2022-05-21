from copy import deepcopy
import os
from PIL import Image
from glob import glob
import pandas as pd
import seaborn as sns
from tqdm import tqdm
import yaml
import docx
import numpy as np
import matplotlib.pyplot as plt


class HealthCheck:
    """
    Dataset health check before training. There are:
        *   Length check: Lengths of images and labels.
        *   Pair check: Check for all images have their label and all labels have their images.
        *   Size Distribution: İmages sizes check (width,height).
        *   Classes Balance: For each classes, how many objects they have and relations between classes according to objects lengths.
        *   Heat map: For each classes, The area occupied by their objects according to dataset.
    """

    def __init__(self, data_yaml: str, target_folder: str = "./") -> None:
        """
        Args:
            data_yaml (str): Dataset yaml file like YOLOV5 data.yaml
            target_folder (str, optional): Target folder for save health check results. Defaults to "./".
        """
        self.target_folder = target_folder
        if not os.path.exists(target_folder):
            os.makedirs(target_folder)
            
        self.data = open(data_yaml, "r").read()
        self.data = yaml.safe_load(self.data)

        self.images = glob(self.data["train"]+"/**")

        label_path = self.data["train"].rsplit("/")
        label_path[-1] = "labels"
        label_path = "/".join(label_path)
        self.labels = glob(label_path+"/*.txt")

        self.save = target_folder+"/"+self.data["train"].rsplit("/")[-2]
        self.counter_df = pd.DataFrame({"classes": []})

    def lengthCheck(self) -> tuple:
        return (len(self.images), len(self.labels))

    def pairCheck(self) -> tuple:
        im_missing = []
        lb_missing = []
        
        images_names = [x.rsplit("/")[-1].rsplit(".")[-2] for x in self.images]
        labels_names = [x.rsplit("/")[-1].rsplit(".")[-2] for x in self.labels]
        
        for e, im in enumerate(tqdm(images_names, "pair check: ")):
            if im not in labels_names:
                im_missing.append(self.images[e])
                
        for e, lb in enumerate(tqdm(labels_names, "pair check:")):
            if lb not in images_names:
                lb_missing.append(self.labels[e])
                
        return (im_missing, lb_missing)

    def sizeDistribution(self) -> None:
        fig_save = f"{self.save}_size.png"
        csv_save = f"{self.save}_size.csv"
        size_df = pd.DataFrame({"image": [], "width": [], "height": []})
        
        for e, img in enumerate(tqdm(self.images, "size dist: ")):
            img_size = Image.open(img).size
            size_df.loc[e, "image"] = img
            size_df.loc[e, "width"] = img_size[0]
            size_df.loc[e, "height"] = img_size[1]

        fig = sns.scatterplot(data=size_df, x="width",
                              y="height", sizes=(20, 200), legend="full")
        fig = fig.get_figure()
        fig.savefig(fig_save)
        size_df.to_csv(csv_save)
        plt.close()

    def classBalance(self) -> None:
        fig_save = f"{self.save}_balance.png"
        csv_save = f"{self.save}_balance.csv"

        for txt in tqdm(self.labels, "class balance check: "):
            f = open(txt, "r").read().splitlines()
            for line in f:
                self.counter_df.loc[len(
                    self.counter_df.classes), "classes"] = self.data["names"][int(line.rsplit(" ")[0])]

        fig = sns.histplot(data=self.counter_df, y="classes", hue="classes")
        fig = fig.get_figure()
        fig.savefig(fig_save)
        self.counter_df.to_csv(csv_save)
        plt.close()

    def xywhn2xyxy(self, xywhn, size) -> list:
        xyxy = deepcopy(xywhn)
        xyxy[0] = int(size[0] * (xywhn[0] - xywhn[2] / 2))
        xyxy[1] = int(size[1] * (xywhn[1] - xywhn[3] / 2))
        xyxy[2] = int(size[0] * (xywhn[0] + xywhn[2] / 2))
        xyxy[3] = int(size[1] * (xywhn[1] + xywhn[3] / 2))
        return xyxy

    def heatmap(self, size=(1024, 1024)) -> None:
        hmap_dict = {}
        for key in self.data["names"]:
            hmap_dict[key] = np.zeros(size, np.int16)

        for txt in tqdm(self.labels, "heat map check: "):
            f = open(txt, "r").read().splitlines()
            for line in f:
                line = [float(x) for x in line.rsplit(" ")]
                xyxy = self.xywhn2xyxy(line[1:], size)
                hmap_dict[self.data["names"]
                          [int(line[0])]][xyxy[0]:xyxy[2], xyxy[1]:xyxy[3]] += 1

        for hmap in hmap_dict.keys():
            sns.heatmap(hmap_dict[hmap]).get_figure().savefig(
                f"{self.save}_heatmap_{hmap}.png")
            plt.close()

    def show(self) -> None:
        doc_name = self.save.rsplit("/")[-1]
        doc = docx.Document()
        
        im_length, lb_length = self.lengthCheck()
        im_missing, lb_missing = self.pairCheck()
        self.sizeDistribution()
        self.classBalance()
        self.heatmap()
        
        doc.add_heading(f"{doc_name} Health check", 0)

        doc.add_heading(f"Length Check", 1)
        doc.add_paragraph(
            f"{doc_name} train has {im_length} images and {lb_length} labels.")

        doc.add_heading(f"Missing Check", 1)
        if len(im_missing) == 0 and len(lb_missing) == 0:
            doc.add_paragraph(
                "There is no missing data. All images have their labels.")
        if len(im_missing) > 0:
            doc.add_paragraph("These images have no their labels:")
            for d in im_missing:
                doc.add_paragraph(d, style="List Bullet")
        if len(lb_missing) > 0:
            doc.add_paragraph("These labels have no their images:")
            for d in lb_missing:
                doc.add_paragraph(d, style="List Bullet")

        doc.add_heading(f"Size Distribution", 1)
        doc.add_picture(f"{self.save}_size.png", width=docx.shared.Inches(
            5), height=docx.shared.Inches(5))
        doc.add_paragraph(f"Figure 1: All image sizes in {doc_name} train.")
        doc.add_paragraph()

        doc.add_heading(f"Classes Balance", 1)
        doc.add_picture(f"{self.save}_balance.png", width=docx.shared.Inches(
            5), height=docx.shared.Inches(5))
        doc.add_paragraph(f"Figure 2: Class in {doc_name} train.")
        doc.add_paragraph()

        plus = 3
        for name in self.data["names"]:
            doc.add_heading(f"Heat map {name} ", 1)
            doc.add_picture(f"{self.save}_heatmap_{name}.png", width=docx.shared.Inches(
                5), height=docx.shared.Inches(5))
            doc.add_paragraph(f"Figure {plus}: Heatmap {doc_name} train.")
            doc.add_paragraph()
            plus += 1

        doc.add_paragraph(f"For details: {self.target_folder}")
        doc.save(f"{self.save}_health_check.docx")


if __name__ == "__main__":
    sns.set_style('darkgrid')
    data_yaml = "C:/Users/Default.DESKTOP-HLSC8QP/Desktop/pistv3/data.yaml"
    target_folder = "C:/Users/Default.DESKTOP-HLSC8QP/Desktop/pistv3/healthcheck"
    HC = HealthCheck(data_yaml, target_folder)
    HC.show()
